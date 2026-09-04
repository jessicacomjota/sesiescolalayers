import hashlib
from io import BytesIO
from datetime import datetime

import pandas as pd
import plotly.express as px
import streamlit as st

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage
)


# ============================================================
# CONFIGURAÇÃO
# ============================================================

st.set_page_config(
    page_title="SGA Layers",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ============================================================
# CSS
# ============================================================

st.markdown(
    """
    <style>

    .main {
        background-color: #f5f7fa;
    }

    .block-container {
        padding-top: 1.5rem;
        padding-bottom: 2rem;
    }

    .titulo-principal {
        font-size: 32px;
        font-weight: 700;
        color: #1f2937;
        margin-bottom: 5px;
    }

    .subtitulo {
        font-size: 16px;
        color: #6b7280;
        margin-bottom: 25px;
    }

    .card {
        background-color: white;
        border-radius: 12px;
        padding: 18px;
        border: 1px solid #e5e7eb;
        box-shadow: 0 1px 5px rgba(0,0,0,0.08);
        margin-bottom: 15px;
    }

    .card-titulo {
        font-size: 13px;
        color: #6b7280;
        margin-bottom: 5px;
    }

    .card-valor {
        font-size: 28px;
        font-weight: 700;
        color: #111827;
    }

    .titulo-secao {
        font-size: 24px;
        font-weight: 700;
        color: #1f2937;
        margin-top: 10px;
        margin-bottom: 20px;
    }

    </style>
    """,
    unsafe_allow_html=True
)


# ============================================================
# CONSTANTES
# ============================================================

COLUNAS_ESPERADAS = [
    "Número do chamado",
    "Atendente",
    "Nome do solicitante",
    "Email do solicitante",
    "Nome do canal",
    "Data de criação",
    "Última atualização",
    "Assunto do chamado",
    "Histórico de Assuntos do chamado",
    "Fonte do chamado",
    "Status do chamado",
    "Tags",
    "Avaliação",
    "Comentário"
]

STATUS_PENDENTES = [
    "Novo",
    "Aberto",
    "Aguardando"
]


# ============================================================
# AUTENTICAÇÃO
# ============================================================

def hash_senha(senha):
    return hashlib.sha256(
        senha.encode("utf-8")
    ).hexdigest()


def verificar_login(usuario, senha):
    try:
        usuarios = st.secrets["usuarios"]

        if usuario not in usuarios:
            return False

        senha_hash = hash_senha(senha)
        senha_configurada = str(usuarios[usuario])

        # Aceita senha em texto ou hash SHA256
        if senha == senha_configurada:
            return True

        if senha_hash == senha_configurada:
            return True

        return False

    except Exception:
        return False


def tela_login():

    st.markdown(
        '<div class="titulo-principal">SGA Layers</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="subtitulo">'
        'Sistema de Gestão de Atendimentos Layers'
        '</div>',
        unsafe_allow_html=True
    )

    col1, col2, col3 = st.columns([1, 1, 1])

    with col2:

        st.markdown("### 🔐 Acesso ao sistema")

        usuario = st.text_input(
            "Usuário",
            key="login_usuario"
        )

        senha = st.text_input(
            "Senha",
            type="password",
            key="login_senha"
        )

        entrar = st.button(
            "Entrar",
            use_container_width=True
        )

        if entrar:

            if verificar_login(usuario, senha):

                st.session_state["autenticado"] = True
                st.session_state["usuario"] = usuario

                st.rerun()

            else:

                st.error(
                    "Usuário ou senha inválidos."
                )

    return False


# ============================================================
# LEITURA DO CSV
# ============================================================

@st.cache_data
def ler_csv(arquivo):

    erros = []

    for encoding in [
        "utf-8-sig",
        "utf-8",
        "cp1252",
        "latin1"
    ]:

        for separador in [
            ",",
            ";",
            "\t"
        ]:

            try:

                arquivo.seek(0)

                df = pd.read_csv(
                    arquivo,
                    sep=separador,
                    encoding=encoding,
                    dtype=str
                )

                if len(df.columns) > 1:

                    df.columns = [
                        str(col).strip()
                        for col in df.columns
                    ]

                    return df

            except Exception as erro:

                erros.append(
                    f"{encoding} / {repr(separador)}: {erro}"
                )

    raise ValueError(
        "Não foi possível ler o arquivo CSV."
    )


# ============================================================
# LIMPEZA DE TEXTO
# ============================================================

def limpar_texto(valor):

    if pd.isna(valor):
        return ""

    valor = str(valor).strip()

    if valor in ["-", "nan", "None"]:
        return ""

    return valor


# ============================================================
# TRATAMENTO DOS DADOS
# ============================================================

@st.cache_data
def tratar_dados(df):

    df = df.copy()

    # --------------------------------------------------------
    # Garantir colunas
    # --------------------------------------------------------

    for coluna in COLUNAS_ESPERADAS:

        if coluna not in df.columns:

            df[coluna] = ""

    # --------------------------------------------------------
    # Limpeza geral
    # --------------------------------------------------------

    for coluna in COLUNAS_ESPERADAS:

        df[coluna] = df[coluna].apply(
            limpar_texto
        )

    # --------------------------------------------------------
    # Campos padronizados
    # --------------------------------------------------------

    df["Atendente"] = df["Atendente"].replace(
        "",
        "Não atribuído"
    )

    df["Tags"] = df["Tags"].replace(
        "",
        "Não informado"
    )

    df["Avaliação"] = df["Avaliação"].replace(
        "",
        "Não informado"
    )

    df["Comentário"] = df["Comentário"].replace(
        "",
        "Não informado"
    )

    df["Fonte do chamado"] = df[
        "Fonte do chamado"
    ].replace(
        "",
        "Não informado"
    )

    # --------------------------------------------------------
    # Histórico
    # --------------------------------------------------------

    df["Histórico de Assuntos do chamado"] = (
        df["Histórico de Assuntos do chamado"]
        .str.rstrip(";")
        .str.strip()
    )

    # --------------------------------------------------------
    # Número
    # --------------------------------------------------------

    df["Número do chamado"] = (
        df["Número do chamado"]
        .astype(str)
        .str.strip()
    )

    # --------------------------------------------------------
    # Datas
    # --------------------------------------------------------

    df["Data de criação"] = pd.to_datetime(
        df["Data de criação"],
        dayfirst=True,
        errors="coerce"
    )

    df["Última atualização"] = pd.to_datetime(
        df["Última atualização"],
        dayfirst=True,
        errors="coerce"
    )

    # --------------------------------------------------------
    # Tempo de atendimento
    # --------------------------------------------------------

    df["Tempo de atendimento"] = (
        df["Última atualização"]
        - df["Data de criação"]
    )

    df["Tempo de atendimento_min"] = (
        df["Tempo de atendimento"]
        .dt.total_seconds()
        / 60
    )

    # Tempos negativos são considerados inválidos
    df.loc[
        df["Tempo de atendimento_min"] < 0,
        "Tempo de atendimento_min"
    ] = pd.NA

    df["Tempo de atendimento_horas"] = (
        df["Tempo de atendimento_min"] / 60
    )

    # --------------------------------------------------------
    # Data e mês
    # --------------------------------------------------------

    df["Data"] = df["Data de criação"].dt.date

    df["Mês"] = (
        df["Data de criação"]
        .dt.to_period("M")
        .astype(str)
    )

    return df


# ============================================================
# FORMATAÇÃO DE TEMPO
# ============================================================

def formatar_tempo(minutos):

    if pd.isna(minutos):
        return "Não informado"

    minutos = float(minutos)

    if minutos < 60:
        return f"{minutos:.0f} min"

    horas = minutos / 60

    if horas < 24:
        return f"{horas:.1f} h"

    dias = horas / 24

    return f"{dias:.1f} dias"


# ============================================================
# CARD
# ============================================================

def card(titulo, valor):
    def card(titulo, valor):
        st.metric(
            label=titulo,
            value=valor
        )

# ============================================================
# FILTROS GLOBAIS
# ============================================================

def aplicar_filtros_globais(df):

    st.sidebar.markdown("## 🔎 Filtros")

    if df.empty:

        st.session_state["filtros_aplicados"] = {
            "periodo": "Sem dados",
            "canais": "Todos",
            "atendentes": "Todos"
        }

        return df

    # --------------------------------------------------------
    # Período
    # --------------------------------------------------------

    data_min = df["Data de criação"].min()
    data_max = df["Data de criação"].max()

    if pd.isna(data_min) or pd.isna(data_max):

        data_inicial = None
        data_final = None

    else:

        data_inicial = data_min.date()
        data_final = data_max.date()

    periodo = st.sidebar.date_input(
        "Período",
        value=(
            data_inicial,
            data_final
        ) if data_inicial and data_final else None
    )

    # --------------------------------------------------------
    # Canal
    # --------------------------------------------------------

    canais = sorted(
        df["Nome do canal"]
        .dropna()
        .astype(str)
        .unique()
        .tolist()
    )

    canais_selecionados = st.sidebar.multiselect(
        "Canal",
        canais,
        default=canais
    )

    # --------------------------------------------------------
    # Atendente
    # --------------------------------------------------------

    atendentes = sorted(
        df["Atendente"]
        .dropna()
        .astype(str)
        .unique()
        .tolist()
    )

    atendentes_selecionados = st.sidebar.multiselect(
        "Atendente",
        atendentes,
        default=atendentes
    )

    filtrado = df.copy()

    # --------------------------------------------------------
    # Aplicar período
    # --------------------------------------------------------

    if isinstance(periodo, tuple):

        if len(periodo) == 2:

            inicio, fim = periodo

            if inicio and fim:

                filtrado = filtrado[
                    (
                        filtrado["Data de criação"].dt.date
                        >= inicio
                    )
                    &
                    (
                        filtrado["Data de criação"].dt.date
                        <= fim
                    )
                ]

    # --------------------------------------------------------
    # Aplicar canal
    # --------------------------------------------------------

    if canais_selecionados:

        filtrado = filtrado[
            filtrado["Nome do canal"].isin(
                canais_selecionados
            )
        ]

    # --------------------------------------------------------
    # Aplicar atendente
    # --------------------------------------------------------

    if atendentes_selecionados:

        filtrado = filtrado[
            filtrado["Atendente"].isin(
                atendentes_selecionados
            )
        ]

    # --------------------------------------------------------
    # Guardar filtros
    # --------------------------------------------------------

    if (
        isinstance(periodo, tuple)
        and len(periodo) == 2
        and periodo[0]
        and periodo[1]
    ):

        periodo_texto = (
            f"{periodo[0].strftime('%d/%m/%Y')} "
            f"a "
            f"{periodo[1].strftime('%d/%m/%Y')}"
        )

    else:

        periodo_texto = "Todos os períodos"

    if canais_selecionados == canais:
        canais_texto = "Todos"

    elif canais_selecionados:
        canais_texto = ", ".join(
            canais_selecionados
        )

    else:
        canais_texto = "Nenhum"

    if atendentes_selecionados == atendentes:
        atendentes_texto = "Todos"

    elif atendentes_selecionados:
        atendentes_texto = ", ".join(
            atendentes_selecionados
        )

    else:
        atendentes_texto = "Nenhum"

    st.session_state["filtros_aplicados"] = {
        "periodo": periodo_texto,
        "canais": canais_texto,
        "atendentes": atendentes_texto
    }

    # --------------------------------------------------------
    # Resumo
    # --------------------------------------------------------

    st.sidebar.markdown("---")

    st.sidebar.metric(
        "Chamados encontrados",
        len(filtrado)
    )

    return filtrado


# ============================================================
# DESCRIÇÃO DOS FILTROS
# ============================================================

def obter_filtros_texto():

    filtros = st.session_state.get(
        "filtros_aplicados",
        {}
    )

    return {
        "Período": filtros.get(
            "periodo",
            "Todos"
        ),
        "Canal": filtros.get(
            "canais",
            "Todos"
        ),
        "Atendente": filtros.get(
            "atendentes",
            "Todos"
        )
    }


# ============================================================
# TABELA QUANTITATIVA
# ============================================================

def tabela_quantitativa(
    df,
    coluna,
    nome_coluna=None
):

    if df.empty or coluna not in df.columns:

        return pd.DataFrame()

    nome_coluna = nome_coluna or coluna

    tabela = (
        df[coluna]
        .fillna("Não informado")
        .astype(str)
        .value_counts()
        .reset_index()
    )

    tabela.columns = [
        nome_coluna,
        "Quantidade"
    ]

    total = tabela["Quantidade"].sum()

    if total > 0:

        tabela["Percentual"] = (
            tabela["Quantidade"]
            / total
            * 100
        ).round(1)

    else:

        tabela["Percentual"] = 0

    return tabela


# ============================================================
# GRÁFICOS
# ============================================================

def grafico_status(df):

    tabela = tabela_quantitativa(
        df,
        "Status do chamado",
        "Status"
    )

    if tabela.empty:
        return None

    fig = px.pie(
        tabela,
        names="Status",
        values="Quantidade",
        hole=0.45,
        title="Chamados por status"
    )

    fig.update_traces(
        textinfo="label+percent"
    )

    fig.update_layout(
        height=500,
        margin=dict(
            l=20,
            r=20,
            t=60,
            b=20
        )
    )

    return fig


def grafico_canais(df):

    tabela = tabela_quantitativa(
        df,
        "Nome do canal",
        "Canal"
    )

    if tabela.empty:
        return None

    fig = px.bar(
        tabela,
        x="Canal",
        y="Quantidade",
        text="Quantidade",
        title="Chamados por canal"
    )

    fig.update_layout(
        xaxis_tickangle=-45,
        height=500
    )

    return fig


def grafico_atendentes(df):

    tabela = tabela_quantitativa(
        df,
        "Atendente"
    )

    if tabela.empty:
        return None

    fig = px.bar(
        tabela,
        x="Atendente",
        y="Quantidade",
        text="Quantidade",
        title="Chamados por atendente"
    )

    fig.update_layout(
        xaxis_tickangle=-45,
        height=500
    )

    return fig


def grafico_assuntos(df):

    tabela = tabela_quantitativa(
        df,
        "Assunto do chamado",
        "Assunto"
    ).head(15)

    if tabela.empty:
        return None

    fig = px.bar(
        tabela.sort_values(
            "Quantidade"
        ),
        x="Quantidade",
        y="Assunto",
        orientation="h",
        text="Quantidade",
        title="Principais assuntos"
    )

    fig.update_layout(
        height=600
    )

    return fig


def grafico_mensal(df):

    if df.empty:
        return None

    tabela = (
        df.dropna(
            subset=["Data de criação"]
        )
        .groupby(
            "Mês"
        )
        .size()
        .reset_index(
            name="Quantidade"
        )
    )

    if tabela.empty:
        return None

    fig = px.line(
        tabela,
        x="Mês",
        y="Quantidade",
        markers=True,
        text="Quantidade",
        title="Evolução mensal dos chamados"
    )

    fig.update_traces(
        textposition="top center"
    )

    fig.update_layout(
        height=500
    )

    return fig


def grafico_tempo_atendente(df):

    if df.empty:
        return None

    tabela = (
        df.groupby(
            "Atendente",
            dropna=False
        )["Tempo de atendimento_min"]
        .mean()
        .reset_index()
    )

    tabela["Tempo médio"] = (
        tabela["Tempo de atendimento_min"]
        / 60
    ).round(2)

    tabela = tabela.sort_values(
        "Tempo médio"
    )

    if tabela.empty:
        return None

    fig = px.bar(
        tabela,
        x="Tempo médio",
        y="Atendente",
        orientation="h",
        text="Tempo médio",
        title="Tempo médio de atendimento por atendente",
        labels={
            "Tempo médio": "Horas"
        }
    )

    fig.update_layout(
        height=600
    )

    return fig


def grafico_tempo_fonte(df):

    if df.empty:
        return None

    tabela = (
        df.groupby(
            "Fonte do chamado",
            dropna=False
        )["Tempo de atendimento_min"]
        .mean()
        .reset_index()
    )

    tabela["Tempo médio"] = (
        tabela["Tempo de atendimento_min"]
        / 60
    ).round(2)

    tabela = tabela.sort_values(
        "Tempo médio"
    )

    if tabela.empty:
        return None

    fig = px.bar(
        tabela,
        x="Fonte do chamado",
        y="Tempo médio",
        text="Tempo médio",
        title="Tempo médio por fonte",
        labels={
            "Tempo médio": "Horas"
        }
    )

    fig.update_layout(
        xaxis_tickangle=-45,
        height=500
    )

    return fig


def grafico_status_atendente(df):

    if df.empty:
        return None

    tabela = (
        df.groupby(
            [
                "Atendente",
                "Status do chamado"
            ]
        )
        .size()
        .reset_index(
            name="Quantidade"
        )
    )

    fig = px.bar(
        tabela,
        x="Atendente",
        y="Quantidade",
        color="Status do chamado",
        barmode="stack",
        title="Status por atendente"
    )

    fig.update_layout(
        xaxis_tickangle=-45,
        height=600
    )

    return fig


def grafico_status_canal(df):

    if df.empty:
        return None

    tabela = (
        df.groupby(
            [
                "Nome do canal",
                "Status do chamado"
            ]
        )
        .size()
        .reset_index(
            name="Quantidade"
        )
    )

    fig = px.bar(
        tabela,
        x="Nome do canal",
        y="Quantidade",
        color="Status do chamado",
        barmode="stack",
        title="Status por canal"
    )

    fig.update_layout(
        xaxis_tickangle=-45,
        height=600
    )

    return fig


def grafico_dia_semana(df):

    if df.empty:
        return None

    dados = df.dropna(
        subset=["Data de criação"]
    ).copy()

    if dados.empty:
        return None

    ordem = [
        "segunda-feira",
        "terça-feira",
        "quarta-feira",
        "quinta-feira",
        "sexta-feira",
        "sábado",
        "domingo"
    ]

    dados["Dia da semana"] = (
        dados["Data de criação"]
        .dt.day_name("pt_BR")
        .str.lower()
    )

    tabela = (
        dados["Dia da semana"]
        .value_counts()
        .reindex(
            ordem,
            fill_value=0
        )
        .reset_index()
    )

    tabela.columns = [
        "Dia da semana",
        "Quantidade"
    ]

    fig = px.bar(
        tabela,
        x="Dia da semana",
        y="Quantidade",
        text="Quantidade",
        title="Chamados por dia da semana"
    )

    fig.update_layout(
        height=500
    )

    return fig


def grafico_hora(df):

    if df.empty:
        return None

    dados = df.dropna(
        subset=["Data de criação"]
    ).copy()

    if dados.empty:
        return None

    dados["Hora"] = (
        dados["Data de criação"]
        .dt.hour
    )

    tabela = (
        dados["Hora"]
        .value_counts()
        .sort_index()
        .reset_index()
    )

    tabela.columns = [
        "Hora",
        "Quantidade"
    ]

    tabela["Horário"] = (
        tabela["Hora"]
        .astype(str)
        .str.zfill(2)
        + ":00"
    )

    fig = px.bar(
        tabela,
        x="Horário",
        y="Quantidade",
        text="Quantidade",
        title="Chamados por horário"
    )

    fig.update_layout(
        height=500
    )

    return fig


# ============================================================
# PREPARAR EXPORTAÇÃO
# ============================================================

def preparar_exportacao(df):

    exportar = df.copy()

    if "Tempo de atendimento_min" in exportar.columns:

        exportar["Tempo de atendimento"] = (
            exportar["Tempo de atendimento_min"]
            .apply(formatar_tempo)
        )

    if "Data de criação" in exportar.columns:

        exportar["Data de criação"] = (
            exportar["Data de criação"]
            .dt.strftime("%d/%m/%Y %H:%M")
        )

    if "Última atualização" in exportar.columns:

        exportar["Última atualização"] = (
            exportar["Última atualização"]
            .dt.strftime("%d/%m/%Y %H:%M")
        )

    colunas_excluir = [
        "Tempo de atendimento_min",
        "Tempo de atendimento_horas",
        "Data",
        "Mês"
    ]

    exportar = exportar.drop(
        columns=[
            coluna
            for coluna in colunas_excluir
            if coluna in exportar.columns
        ],
        errors="ignore"
    )

    return exportar


# ============================================================
# EXCEL
# ============================================================

def gerar_excel(df):

    output = BytesIO()

    with pd.ExcelWriter(
        output,
        engine="openpyxl"
    ) as writer:

        preparar_exportacao(df).to_excel(
            writer,
            index=False,
            sheet_name="Chamados"
        )

    output.seek(0)

    return output.getvalue()


def gerar_excel_geral(df):

    output = BytesIO()

    with pd.ExcelWriter(
        output,
        engine="openpyxl"
    ) as writer:

        # ----------------------------------------------------
        # Resumo
        # ----------------------------------------------------

        total = len(df)

        resumo = pd.DataFrame({
            "Indicador": [
                "Total de chamados",
                "Novos",
                "Abertos",
                "Aguardando",
                "Resolvidos",
                "Tempo médio de atendimento",
                "Taxa de resolução"
            ],
            "Valor": [
                total,
                len(
                    df[
                        df["Status do chamado"]
                        == "Novo"
                    ]
                ),
                len(
                    df[
                        df["Status do chamado"]
                        == "Aberto"
                    ]
                ),
                len(
                    df[
                        df["Status do chamado"]
                        == "Aguardando"
                    ]
                ),
                len(
                    df[
                        df["Status do chamado"]
                        == "Resolvido"
                    ]
                ),
                formatar_tempo(
                    df["Tempo de atendimento_min"]
                    .mean()
                ),
                (
                    len(
                        df[
                            df["Status do chamado"]
                            == "Resolvido"
                        ]
                    )
                    / total
                    * 100
                    if total > 0
                    else 0
                )
            ]
        })

        resumo.to_excel(
            writer,
            index=False,
            sheet_name="Resumo"
        )

        # ----------------------------------------------------
        # Status
        # ----------------------------------------------------

        tabela_quantitativa(
            df,
            "Status do chamado",
            "Status"
        ).to_excel(
            writer,
            index=False,
            sheet_name="Status"
        )

        # ----------------------------------------------------
        # Fontes
        # ----------------------------------------------------

        tabela_quantitativa(
            df,
            "Fonte do chamado",
            "Fonte"
        ).to_excel(
            writer,
            index=False,
            sheet_name="Fontes"
        )

        # ----------------------------------------------------
        # Canais
        # ----------------------------------------------------

        tabela_quantitativa(
            df,
            "Nome do canal",
            "Canal"
        ).to_excel(
            writer,
            index=False,
            sheet_name="Canais"
        )

        # ----------------------------------------------------
        # Assuntos
        # ----------------------------------------------------

        tabela_quantitativa(
            df,
            "Assunto do chamado",
            "Assunto"
        ).to_excel(
            writer,
            index=False,
            sheet_name="Assuntos"
        )

        # ----------------------------------------------------
        # Atendentes
        # ----------------------------------------------------

        tabela_quantitativa(
            df,
            "Atendente"
        ).to_excel(
            writer,
            index=False,
            sheet_name="Atendentes"
        )

        # ----------------------------------------------------
        # Tempos
        # ----------------------------------------------------

        if not df.empty:

            tempos = (
                df.groupby(
                    "Atendente"
                )["Tempo de atendimento_min"]
                .agg(
                    [
                        "count",
                        "mean",
                        "median",
                        "min",
                        "max"
                    ]
                )
                .reset_index()
            )

            tempos.columns = [
                "Atendente",
                "Quantidade",
                "Média (min)",
                "Mediana (min)",
                "Mínimo (min)",
                "Máximo (min)"
            ]

        else:

            tempos = pd.DataFrame()

        tempos.to_excel(
            writer,
            index=False,
            sheet_name="Tempos"
        )

        # ----------------------------------------------------
        # Evolução mensal
        # ----------------------------------------------------

        if not df.empty:

            evolucao = (
                df.groupby(
                    "Mês"
                )
                .size()
                .reset_index(
                    name="Quantidade"
                )
            )

        else:

            evolucao = pd.DataFrame()

        evolucao.to_excel(
            writer,
            index=False,
            sheet_name="Evolução Mensal"
        )

        # ----------------------------------------------------
        # Pendências
        # ----------------------------------------------------

        pendencias = df[
            df["Status do chamado"]
            .isin(STATUS_PENDENTES)
        ]

        preparar_exportacao(
            pendencias
        ).to_excel(
            writer,
            index=False,
            sheet_name="Pendências"
        )

        # ----------------------------------------------------
        # Todos os chamados
        # ----------------------------------------------------

        preparar_exportacao(
            df
        ).to_excel(
            writer,
            index=False,
            sheet_name="Chamados"
        )

    output.seek(0)

    return output.getvalue()


# ============================================================
# WORD - FUNÇÕES AUXILIARES
# ============================================================

def configurar_documento():

    doc = Document()

    section = doc.sections[0]

    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    return doc


def adicionar_titulo(
    doc,
    titulo,
    subtitulo=None
):

    p = doc.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(titulo)

    run.bold = True
    run.font.size = Pt(20)

    if subtitulo:

        p2 = doc.add_paragraph()

        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run2 = p2.add_run(
            subtitulo
        )

        run2.font.size = Pt(10)

        run2.italic = True


def adicionar_tabela_word(
    doc,
    df,
    max_linhas=500
):

    if df is None or df.empty:

        doc.add_paragraph(
            "Não há dados para apresentar."
        )

        return

    tabela_df = df.head(
        max_linhas
    ).copy()

    tabela = doc.add_table(
        rows=1,
        cols=len(tabela_df.columns)
    )

    tabela.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    tabela.style = "Table Grid"

    # Cabeçalho
    for i, coluna in enumerate(
        tabela_df.columns
    ):

        tabela.rows[0].cells[i].text = str(
            coluna
        )

    # Dados
    for _, linha in tabela_df.iterrows():

        cells = tabela.add_row().cells

        for i, valor in enumerate(linha):

            cells[i].text = str(
                valor
            )

    if len(df) > max_linhas:

        doc.add_paragraph(
            f"Exibidas {max_linhas} linhas. "
            "A relação completa está disponível "
            "na exportação Excel."
        )


def inserir_grafico_word(
    doc,
    fig,
    titulo=None
):

    if fig is None:
        return

    try:

        imagem = fig.to_image(
            format="png",
            width=1200,
            height=650,
            scale=2
        )

        if titulo:

            p = doc.add_paragraph()

            run = p.add_run(
                titulo
            )

            run.bold = True

        doc.add_picture(
            BytesIO(imagem),
            width=Inches(7.0)
        )

        doc.paragraphs[-1].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    except Exception as erro:

        doc.add_paragraph(
            "Gráfico não pôde ser inserido. "
            "Verifique se o pacote Kaleido está instalado."
        )


# ============================================================
# WORD - CHAMADOS
# ============================================================

def gerar_docx_chamados(df):

    doc = configurar_documento()

    adicionar_titulo(
        doc,
        "SGA Layers — Chamados",
        f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )

    doc.add_paragraph(
        f"Quantidade de chamados analisados: {len(df)}"
    )

    doc.add_paragraph(
        "Filtros aplicados:"
    )

    for chave, valor in obter_filtros_texto().items():

        doc.add_paragraph(
            f"{chave}: {valor}",
            style="List Bullet"
        )

    doc.add_paragraph()

    tabela = preparar_exportacao(
        df
    )

    adicionar_tabela_word(
        doc,
        tabela,
        max_linhas=500
    )

    output = BytesIO()

    doc.save(output)

    output.seek(0)

    return output.getvalue()


# ============================================================
# WORD - RELATÓRIO GERAL
# ============================================================

def gerar_docx_geral(df):

    doc = configurar_documento()

    adicionar_titulo(
        doc,
        "SGA Layers",
        "Relatório Geral de Atendimentos"
    )

    doc.add_paragraph(
        f"Relatório gerado em "
        f"{datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )

    doc.add_heading(
        "Filtros aplicados",
        level=1
    )

    for chave, valor in obter_filtros_texto().items():

        doc.add_paragraph(
            f"{chave}: {valor}"
        )

    doc.add_paragraph(
        f"Chamados analisados: {len(df)}"
    )

    # --------------------------------------------------------
    # Resumo
    # --------------------------------------------------------

    doc.add_heading(
        "1. Resumo geral",
        level=1
    )

    total = len(df)

    resolvidos = len(
        df[
            df["Status do chamado"]
            == "Resolvido"
        ]
    )

    tempo_medio = (
        df["Tempo de atendimento_min"]
        .mean()
    )

    resumo = pd.DataFrame({
        "Indicador": [
            "Total de chamados",
            "Novos",
            "Abertos",
            "Aguardando",
            "Resolvidos",
            "Pendentes",
            "Tempo médio",
            "Taxa de resolução"
        ],
        "Quantidade / Valor": [
            total,
            len(
                df[
                    df["Status do chamado"]
                    == "Novo"
                ]
            ),
            len(
                df[
                    df["Status do chamado"]
                    == "Aberto"
                ]
            ),
            len(
                df[
                    df["Status do chamado"]
                    == "Aguardando"
                ]
            ),
            resolvidos,
            len(
                df[
                    df["Status do chamado"]
                    .isin(STATUS_PENDENTES)
                ]
            ),
            formatar_tempo(
                tempo_medio
            ),
            f"{(resolvidos / total * 100):.1f}%"
            if total > 0
            else "0%"
        ]
    })

    adicionar_tabela_word(
        doc,
        resumo
    )

    # --------------------------------------------------------
    # Status
    # --------------------------------------------------------

    doc.add_heading(
        "2. Chamados por status",
        level=1
    )

    fig = grafico_status(df)

    inserir_grafico_word(
        doc,
        fig
    )

    tabela_status = tabela_quantitativa(
        df,
        "Status do chamado",
        "Status"
    )

    adicionar_tabela_word(
        doc,
        tabela_status
    )

    # --------------------------------------------------------
    # Canal
    # --------------------------------------------------------

    doc.add_heading(
        "3. Chamados por canal",
        level=1
    )

    fig = grafico_canais(df)

    inserir_grafico_word(
        doc,
        fig
    )

    tabela_canal = tabela_quantitativa(
        df,
        "Nome do canal",
        "Canal"
    )

    adicionar_tabela_word(
        doc,
        tabela_canal
    )

    # --------------------------------------------------------
    # Atendente
    # --------------------------------------------------------

    doc.add_heading(
        "4. Chamados por atendente",
        level=1
    )

    fig = grafico_atendentes(df)

    inserir_grafico_word(
        doc,
        fig
    )

    tabela_atendente = tabela_quantitativa(
        df,
        "Atendente"
    )

    adicionar_tabela_word(
        doc,
        tabela_atendente
    )

    # --------------------------------------------------------
    # Assuntos
    # --------------------------------------------------------

    doc.add_heading(
        "5. Principais assuntos",
        level=1
    )

    fig = grafico_assuntos(df)

    inserir_grafico_word(
        doc,
        fig
    )

    tabela_assuntos = tabela_quantitativa(
        df,
        "Assunto do chamado",
        "Assunto"
    ).head(15)

    adicionar_tabela_word(
        doc,
        tabela_assuntos
    )

    # --------------------------------------------------------
    # Evolução mensal
    # --------------------------------------------------------

    doc.add_heading(
        "6. Evolução mensal",
        level=1
    )

    fig = grafico_mensal(df)

    inserir_grafico_word(
        doc,
        fig
    )

    if not df.empty:

        evolucao = (
            df.groupby("Mês")
            .size()
            .reset_index(
                name="Quantidade"
            )
        )

    else:

        evolucao = pd.DataFrame()

    adicionar_tabela_word(
        doc,
        evolucao
    )

    # --------------------------------------------------------
    # Tempo por atendente
    # --------------------------------------------------------

    doc.add_heading(
        "7. Tempo médio por atendente",
        level=1
    )

    fig = grafico_tempo_atendente(df)

    inserir_grafico_word(
        doc,
        fig
    )

    if not df.empty:

        tempo_atendente = (
            df.groupby(
                "Atendente"
            )["Tempo de atendimento_min"]
            .mean()
            .reset_index()
        )

        tempo_atendente[
            "Tempo médio"
        ] = tempo_atendente[
            "Tempo de atendimento_min"
        ].apply(
            formatar_tempo
        )

        tempo_atendente = (
            tempo_atendente
            .drop(
                columns=[
                    "Tempo de atendimento_min"
                ]
            )
        )

    else:

        tempo_atendente = pd.DataFrame()

    adicionar_tabela_word(
        doc,
        tempo_atendente
    )

    # --------------------------------------------------------
    # Status por atendente
    # --------------------------------------------------------

    doc.add_heading(
        "8. Status por atendente",
        level=1
    )

    fig = grafico_status_atendente(df)

    inserir_grafico_word(
        doc,
        fig
    )

    if not df.empty:

        status_atendente = (
            df.groupby(
                [
                    "Atendente",
                    "Status do chamado"
                ]
            )
            .size()
            .reset_index(
                name="Quantidade"
            )
        )

    else:

        status_atendente = pd.DataFrame()

    adicionar_tabela_word(
        doc,
        status_atendente
    )

    # --------------------------------------------------------
    # Status por canal
    # --------------------------------------------------------

    doc.add_heading(
        "9. Status por canal",
        level=1
    )

    fig = grafico_status_canal(df)

    inserir_grafico_word(
        doc,
        fig
    )

    if not df.empty:

        status_canal = (
            df.groupby(
                [
                    "Nome do canal",
                    "Status do chamado"
                ]
            )
            .size()
            .reset_index(
                name="Quantidade"
            )
        )

    else:

        status_canal = pd.DataFrame()

    adicionar_tabela_word(
        doc,
        status_canal
    )

    # --------------------------------------------------------
    # Pendências
    # --------------------------------------------------------

    doc.add_heading(
        "10. Pendências",
        level=1
    )

    pendencias = df[
        df["Status do chamado"]
        .isin(STATUS_PENDENTES)
    ].copy()

    pendencias_resumo = tabela_quantitativa(
        pendencias,
        "Status do chamado",
        "Status"
    )

    adicionar_tabela_word(
        doc,
        pendencias_resumo
    )

    if not pendencias.empty:

        pendencias_export = preparar_exportacao(
            pendencias
        )

        adicionar_tabela_word(
            doc,
            pendencias_export,
            max_linhas=200
        )

    output = BytesIO()

    doc.save(output)

    output.seek(0)

    return output.getvalue()


# ============================================================
# PDF - ESTILOS
# ============================================================

def criar_estilos_pdf():

    styles = getSampleStyleSheet()

    styles.add(
        ParagraphStyle(
            name="TituloSGA",
            parent=styles["Title"],
            fontSize=20,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=15
        )
    )

    styles.add(
        ParagraphStyle(
            name="SubtituloSGA",
            parent=styles["Normal"],
            fontSize=9,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=15
        )
    )

    styles.add(
        ParagraphStyle(
            name="SecaoSGA",
            parent=styles["Heading1"],
            fontSize=14,
            leading=18,
            spaceBefore=10,
            spaceAfter=8
        )
    )

    return styles


# ============================================================
# PDF - TABELA
# ============================================================

def dataframe_para_tabela_pdf(
    df,
    max_linhas=100
):

    if df is None or df.empty:

        return Table(
            [["Não há dados para apresentar."]]
        )

    tabela_df = df.head(
        max_linhas
    ).copy()

    dados = [
        [
            str(coluna)
            for coluna in tabela_df.columns
        ]
    ]

    for _, linha in tabela_df.iterrows():

        dados.append(
            [
                str(valor)
                for valor in linha
            ]
        )

    tabela = Table(
        dados,
        repeatRows=1
    )

    tabela.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0, 0),
                    (-1, 0),
                    colors.HexColor("#1f2937")
                ),
                (
                    "TEXTCOLOR",
                    (0, 0),
                    (-1, 0),
                    colors.white
                ),
                (
                    "FONTNAME",
                    (0, 0),
                    (-1, 0),
                    "Helvetica-Bold"
                ),
                (
                    "FONTSIZE",
                    (0, 0),
                    (-1, -1),
                    7
                ),
                (
                    "GRID",
                    (0, 0),
                    (-1, -1),
                    0.25,
                    colors.grey
                ),
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "MIDDLE"
                ),
                (
                    "ALIGN",
                    (0, 0),
                    (-1, -1),
                    "CENTER"
                ),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [
                        colors.white,
                        colors.HexColor("#f3f4f6")
                    ]
                )
            ]
        )
    )

    return tabela


# ============================================================
# PDF - GRÁFICO
# ============================================================

def gerar_imagem_grafico(fig):

    if fig is None:
        return None

    try:

        imagem = fig.to_image(
            format="png",
            width=1200,
            height=650,
            scale=2
        )

        return BytesIO(imagem)

    except Exception:

        return None


# ============================================================
# PDF - CHAMADOS
# ============================================================

def gerar_pdf_chamados(df):

    output = BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=1 * cm,
        leftMargin=1 * cm,
        topMargin=1 * cm,
        bottomMargin=1 * cm
    )

    styles = criar_estilos_pdf()

    elementos = []

    elementos.append(
        Paragraph(
            "SGA Layers — Chamados",
            styles["TituloSGA"]
        )
    )

    elementos.append(
        Paragraph(
            f"Chamados analisados: {len(df)}",
            styles["SubtituloSGA"]
        )
    )

    elementos.append(
        Paragraph(
            f"Gerado em "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')}",
            styles["SubtituloSGA"]
        )
    )

    elementos.append(
        Paragraph(
            "Filtros aplicados",
            styles["SecaoSGA"]
        )
    )

    filtros = obter_filtros_texto()

    filtro_df = pd.DataFrame(
        {
            "Filtro": list(
                filtros.keys()
            ),
            "Valor": list(
                filtros.values()
            )
        }
    )

    elementos.append(
        dataframe_para_tabela_pdf(
            filtro_df
        )
    )

    elementos.append(
        Spacer(1, 0.4 * cm)
    )

    elementos.append(
        Paragraph(
            "Chamados",
            styles["SecaoSGA"]
        )
    )

    elementos.append(
        dataframe_para_tabela_pdf(
            preparar_exportacao(df),
            max_linhas=100
        )
    )

    doc.build(elementos)

    output.seek(0)

    return output.getvalue()


# ============================================================
# PDF - RELATÓRIO GERAL
# ============================================================

def gerar_pdf_geral(df):

    output = BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=1 * cm,
        leftMargin=1 * cm,
        topMargin=1 * cm,
        bottomMargin=1 * cm
    )

    styles = criar_estilos_pdf()

    elementos = []

    # --------------------------------------------------------
    # Capa
    # --------------------------------------------------------

    elementos.append(
        Spacer(1, 3 * cm)
    )

    elementos.append(
        Paragraph(
            "SGA Layers",
            styles["TituloSGA"]
        )
    )

    elementos.append(
        Paragraph(
            "RELATÓRIO GERAL DE ATENDIMENTOS",
            styles["TituloSGA"]
        )
    )

    elementos.append(
        Paragraph(
            f"Gerado em "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')}",
            styles["SubtituloSGA"]
        )
    )

    elementos.append(
        Spacer(1, 1 * cm)
    )

    elementos.append(
        Paragraph(
            f"Chamados analisados: {len(df)}",
            styles["SubtituloSGA"]
        )
    )

    elementos.append(
        Spacer(1, 4 * cm)
    )

    # --------------------------------------------------------
    # Filtros
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "Filtros aplicados",
            styles["SecaoSGA"]
        )
    )

    filtros = obter_filtros_texto()

    filtro_df = pd.DataFrame(
        {
            "Filtro": list(
                filtros.keys()
            ),
            "Valor": list(
                filtros.values()
            )
        }
    )

    elementos.append(
        dataframe_para_tabela_pdf(
            filtro_df
        )
    )

    elementos.append(
        Spacer(1, 0.5 * cm)
    )

    # --------------------------------------------------------
    # Resumo
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "1. Resumo geral",
            styles["SecaoSGA"]
        )
    )

    total = len(df)

    resolvidos = len(
        df[
            df["Status do chamado"]
            == "Resolvido"
        ]
    )

    resumo = pd.DataFrame({
        "Indicador": [
            "Total",
            "Novos",
            "Abertos",
            "Aguardando",
            "Resolvidos",
            "Pendentes",
            "Tempo médio",
            "Taxa de resolução"
        ],
        "Valor": [
            total,
            len(
                df[
                    df["Status do chamado"]
                    == "Novo"
                ]
            ),
            len(
                df[
                    df["Status do chamado"]
                    == "Aberto"
                ]
            ),
            len(
                df[
                    df["Status do chamado"]
                    == "Aguardando"
                ]
            ),
            resolvidos,
            len(
                df[
                    df["Status do chamado"]
                    .isin(STATUS_PENDENTES)
                ]
            ),
            formatar_tempo(
                df[
                    "Tempo de atendimento_min"
                ].mean()
            ),
            (
                f"{resolvidos / total * 100:.1f}%"
                if total > 0
                else "0%"
            )
        ]
    })

    elementos.append(
        dataframe_para_tabela_pdf(
            resumo
        )
    )

    # --------------------------------------------------------
    # Status
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "2. Chamados por status",
            styles["SecaoSGA"]
        )
    )

    imagem = gerar_imagem_grafico(
        grafico_status(df)
    )

    if imagem:

        elementos.append(
            RLImage(
                imagem,
                width=22 * cm,
                height=12 * cm
            )
        )

    elementos.append(
        dataframe_para_tabela_pdf(
            tabela_quantitativa(
                df,
                "Status do chamado",
                "Status"
            )
        )
    )

    # --------------------------------------------------------
    # Canal
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "3. Chamados por canal",
            styles["SecaoSGA"]
        )
    )

    imagem = gerar_imagem_grafico(
        grafico_canais(df)
    )

    if imagem:

        elementos.append(
            RLImage(
                imagem,
                width=22 * cm,
                height=12 * cm
            )
        )

    elementos.append(
        dataframe_para_tabela_pdf(
            tabela_quantitativa(
                df,
                "Nome do canal",
                "Canal"
            )
        )
    )

    # --------------------------------------------------------
    # Atendente
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "4. Chamados por atendente",
            styles["SecaoSGA"]
        )
    )

    imagem = gerar_imagem_grafico(
        grafico_atendentes(df)
    )

    if imagem:

        elementos.append(
            RLImage(
                imagem,
                width=22 * cm,
                height=12 * cm
            )
        )

    elementos.append(
        dataframe_para_tabela_pdf(
            tabela_quantitativa(
                df,
                "Atendente"
            )
        )
    )

    # --------------------------------------------------------
    # Assuntos
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "5. Principais assuntos",
            styles["SecaoSGA"]
        )
    )

    imagem = gerar_imagem_grafico(
        grafico_assuntos(df)
    )

    if imagem:

        elementos.append(
            RLImage(
                imagem,
                width=22 * cm,
                height=12 * cm
            )
        )

    elementos.append(
        dataframe_para_tabela_pdf(
            tabela_quantitativa(
                df,
                "Assunto do chamado",
                "Assunto"
            ).head(15)
        )
    )

    # --------------------------------------------------------
    # Evolução mensal
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "6. Evolução mensal",
            styles["SecaoSGA"]
        )
    )

    imagem = gerar_imagem_grafico(
        grafico_mensal(df)
    )

    if imagem:

        elementos.append(
            RLImage(
                imagem,
                width=22 * cm,
                height=12 * cm
            )
        )

    if not df.empty:

        evolucao = (
            df.groupby("Mês")
            .size()
            .reset_index(
                name="Quantidade"
            )
        )

    else:

        evolucao = pd.DataFrame()

    elementos.append(
        dataframe_para_tabela_pdf(
            evolucao
        )
    )

    # --------------------------------------------------------
    # Tempo por atendente
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "7. Tempo médio por atendente",
            styles["SecaoSGA"]
        )
    )

    imagem = gerar_imagem_grafico(
        grafico_tempo_atendente(df)
    )

    if imagem:

        elementos.append(
            RLImage(
                imagem,
                width=22 * cm,
                height=12 * cm
            )
        )

    if not df.empty:

        tempo_atendente = (
            df.groupby(
                "Atendente"
            )["Tempo de atendimento_min"]
            .mean()
            .reset_index()
        )

        tempo_atendente[
            "Tempo médio"
        ] = tempo_atendente[
            "Tempo de atendimento_min"
        ].apply(
            formatar_tempo
        )

        tempo_atendente = (
            tempo_atendente
            .drop(
                columns=[
                    "Tempo de atendimento_min"
                ]
            )
        )

    else:

        tempo_atendente = pd.DataFrame()

    elementos.append(
        dataframe_para_tabela_pdf(
            tempo_atendente
        )
    )

    # --------------------------------------------------------
    # Status por atendente
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "8. Status por atendente",
            styles["SecaoSGA"]
        )
    )

    imagem = gerar_imagem_grafico(
        grafico_status_atendente(df)
    )

    if imagem:

        elementos.append(
            RLImage(
                imagem,
                width=22 * cm,
                height=12 * cm
            )
        )

    if not df.empty:

        status_atendente = (
            df.groupby(
                [
                    "Atendente",
                    "Status do chamado"
                ]
            )
            .size()
            .reset_index(
                name="Quantidade"
            )
        )

    else:

        status_atendente = pd.DataFrame()

    elementos.append(
        dataframe_para_tabela_pdf(
            status_atendente
        )
    )

    # --------------------------------------------------------
    # Status por canal
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "9. Status por canal",
            styles["SecaoSGA"]
        )
    )

    imagem = gerar_imagem_grafico(
        grafico_status_canal(df)
    )

    if imagem:

        elementos.append(
            RLImage(
                imagem,
                width=22 * cm,
                height=12 * cm
            )
        )

    if not df.empty:

        status_canal = (
            df.groupby(
                [
                    "Nome do canal",
                    "Status do chamado"
                ]
            )
            .size()
            .reset_index(
                name="Quantidade"
            )
        )

    else:

        status_canal = pd.DataFrame()

    elementos.append(
        dataframe_para_tabela_pdf(
            status_canal
        )
    )

    # --------------------------------------------------------
    # Pendências
    # --------------------------------------------------------

    elementos.append(
        Paragraph(
            "10. Pendências",
            styles["SecaoSGA"]
        )
    )

    pendencias = df[
        df["Status do chamado"]
        .isin(STATUS_PENDENTES)
    ]

    elementos.append(
        dataframe_para_tabela_pdf(
            tabela_quantitativa(
                pendencias,
                "Status do chamado",
                "Status"
            )
        )
    )

    doc.build(elementos)

    output.seek(0)

    return output.getvalue()


# ============================================================
# BOTÕES DE EXPORTAÇÃO
# ============================================================

def botoes_exportacao(
    df,
    nome_arquivo="chamados"
):

    col1, col2, col3 = st.columns(3)

    with col1:

        st.download_button(
            "📊 Excel",
            data=gerar_excel(df),
            file_name=f"{nome_arquivo}.xlsx",
            mime=(
                "application/vnd.openxmlformats-"
                "officedocument.spreadsheetml.sheet"
            ),
            use_container_width=True
        )

    with col2:

        st.download_button(
            "📝 Word",
            data=gerar_docx_chamados(df),
            file_name=f"{nome_arquivo}.docx",
            mime=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),
            use_container_width=True
        )

    with col3:

        st.download_button(
            "📄 PDF",
            data=gerar_pdf_chamados(df),
            file_name=f"{nome_arquivo}.pdf",
            mime="application/pdf",
            use_container_width=True
        )


def botoes_relatorio_geral(df):

    st.markdown(
        "### 📥 Exportar relatório geral"
    )

    col1, col2, col3 = st.columns(3)

    with col1:

        st.download_button(
            "📊 Excel Geral",
            data=gerar_excel_geral(df),
            file_name="relatorio_geral_layers.xlsx",
            mime=(
                "application/vnd.openxmlformats-"
                "officedocument.spreadsheetml.sheet"
            ),
            use_container_width=True
        )

    with col2:

        st.download_button(
            "📝 Word Geral",
            data=gerar_docx_geral(df),
            file_name="relatorio_geral_layers.docx",
            mime=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),
            use_container_width=True
        )

    with col3:

        st.download_button(
            "📄 PDF Geral",
            data=gerar_pdf_geral(df),
            file_name="relatorio_geral_layers.pdf",
            mime="application/pdf",
            use_container_width=True
        )


# ============================================================
# DASHBOARD
# ============================================================

def tela_dashboard(df):

    st.markdown(
        '<div class="titulo-secao">📊 DASHBOARD</div>',
        unsafe_allow_html=True
    )

    if df.empty:

        st.warning(
            "Nenhum chamado encontrado com os filtros selecionados."
        )

        return

    total = len(df)

    novos = len(
        df[
            df["Status do chamado"]
            == "Novo"
        ]
    )

    abertos = len(
        df[
            df["Status do chamado"]
            == "Aberto"
        ]
    )

    aguardando = len(
        df[
            df["Status do chamado"]
            == "Aguardando"
        ]
    )

    resolvidos = len(
        df[
            df["Status do chamado"]
            == "Resolvido"
        ]
    )

    tempo_medio = (
        df["Tempo de atendimento_min"]
        .mean()
    )

    col1, col2, col3, col4, col5, col6 = st.columns(6)

    with col1:
        card("Total", total)

    with col2:
        card("Novos", novos)

    with col3:
        card("Abertos", abertos)

    with col4:
        card("Aguardando", aguardando)

    with col5:
        card("Resolvidos", resolvidos)

    with col6:
        card(
            "Tempo médio",
            formatar_tempo(tempo_medio)
        )

    st.markdown("---")

    col1, col2 = st.columns(2)

    with col1:

        fig = grafico_status(df)

        if fig:
            st.plotly_chart(
                fig,
                use_container_width=True
            )

    with col2:

        fig = grafico_canais(df)

        if fig:
            st.plotly_chart(
                fig,
                use_container_width=True
            )

    fig = grafico_assuntos(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    st.markdown(
        "### 📥 Exportação"
    )

    botoes_exportacao(
        df,
        "dashboard_layers"
    )


# ============================================================
# CHAMADOS
# ============================================================

def tela_chamados(df):

    st.markdown(
        '<div class="titulo-secao">📋 CHAMADOS</div>',
        unsafe_allow_html=True
    )

    if df.empty:

        st.warning(
            "Nenhum chamado encontrado."
        )

        return

    # --------------------------------------------------------
    # FILTROS
    # --------------------------------------------------------

    col1, col2, col3, col4 = st.columns(4)

    # --------------------------------------------------------
    # FILTRO STATUS
    # --------------------------------------------------------

    with col1:

        status_opcoes = [
            "Todos"
        ] + sorted(
            df["Status do chamado"]
            .dropna()
            .astype(str)
            .unique()
            .tolist()
        )

        status_filtro = st.selectbox(
            "Status",
            status_opcoes
        )

    # --------------------------------------------------------
    # FILTRO ASSUNTO
    # --------------------------------------------------------

    with col2:

        assuntos = [
            "Todos"
        ] + sorted(
            df["Assunto do chamado"]
            .dropna()
            .astype(str)
            .unique()
            .tolist()
        )

        assunto_filtro = st.selectbox(
            "Assunto",
            assuntos
        )

    # --------------------------------------------------------
    # FILTRO CANAL
    # --------------------------------------------------------

    with col3:

        canais = [
            "Todos"
        ] + sorted(
            df["Nome do canal"]
            .dropna()
            .astype(str)
            .unique()
            .tolist()
        )

        canal_filtro = st.selectbox(
            "Canal",
            canais
        )

    # --------------------------------------------------------
    # BUSCA
    # --------------------------------------------------------

    with col4:

        busca = st.text_input(
            "🔎 Buscar chamado",
            placeholder=(
                "Número, solicitante, "
                "e-mail ou assunto..."
            )
        )

    # --------------------------------------------------------
    # APLICAÇÃO DOS FILTROS
    # --------------------------------------------------------

    resultado = df.copy()

    # Filtro por status
    if status_filtro != "Todos":

        resultado = resultado[
            resultado["Status do chamado"].astype(str)
            == status_filtro
        ]

    # Filtro por assunto
    if assunto_filtro != "Todos":

        resultado = resultado[
            resultado["Assunto do chamado"].astype(str)
            == assunto_filtro
        ]

    # Filtro por canal
    if canal_filtro != "Todos":

        resultado = resultado[
            resultado["Nome do canal"].astype(str)
            == canal_filtro
        ]

    # Busca geral
    if busca:

        busca = busca.lower().strip()

        colunas_busca = [
            "Número do chamado",
            "Nome do solicitante",
            "Email do solicitante",
            "Assunto do chamado"
        ]

        colunas_busca = [
            coluna
            for coluna in colunas_busca
            if coluna in resultado.columns
        ]

        if colunas_busca:

            mascara = (
                resultado[colunas_busca]
                .fillna("")
                .astype(str)
                .apply(
                    lambda coluna:
                    coluna.str.lower().str.contains(
                        busca,
                        na=False,
                        regex=False
                    )
                )
                .any(axis=1)
            )

            resultado = resultado[
                mascara
            ]

    # --------------------------------------------------------
    # TOTAL DE CHAMADOS ENCONTRADOS
    # --------------------------------------------------------

    st.metric(
        "Chamados encontrados",
        len(resultado)
    )

    # --------------------------------------------------------
    # TABELA
    # --------------------------------------------------------

    tabela = preparar_exportacao(
        resultado
    )

    st.dataframe(
        tabela,
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # EXPORTAÇÃO
    # --------------------------------------------------------

    st.markdown(
        "### 📥 Exportação"
    )

    botoes_exportacao(
        resultado,
        "chamados_filtrados_layers"
    )
    # --------------------------------------------------------
    # TABELA
    # --------------------------------------------------------

    tabela = preparar_exportacao(
        resultado
    )

    st.dataframe(
        tabela,
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # EXPORTAÇÃO
    # --------------------------------------------------------

    st.markdown(
        "### 📥 Exportação"
    )

    botoes_exportacao(
        resultado,
        "chamados_filtrados_layers"
    )


# ============================================================
# DETALHAMENTO
# ============================================================

def tela_detalhamento(df):

    st.markdown(
        '<div class="titulo-secao">🔍 DETALHAMENTO</div>',
        unsafe_allow_html=True
    )

    if df.empty:

        st.warning(
            "Nenhum chamado disponível para detalhamento."
        )

        return

    numeros = (
        df["Número do chamado"]
        .dropna()
        .astype(str)
        .tolist()
    )

    numero = st.selectbox(
        "Selecione o chamado",
        numeros
    )

    chamado = df[
        df["Número do chamado"].astype(str)
        == str(numero)
    ]

    if chamado.empty:
        return

    linha = chamado.iloc[0]

    # --------------------------------------------------------
    # INFORMAÇÕES PRINCIPAIS
    # --------------------------------------------------------

    col1, col2, col3 = st.columns(3)

    with col1:

        st.markdown("### 👤 Solicitante")

        st.write(
            linha["Nome do solicitante"]
        )

        st.write(
            linha["Email do solicitante"]
        )

    with col2:

        st.markdown("### 👨‍💼 Atendimento")

        st.write(
            linha["Atendente"]
        )

        st.write(
            linha["Status do chamado"]
        )

    with col3:

        st.markdown("### 🏷️ Classificação")

        st.write(
            linha["Nome do canal"]
        )

        st.write(
            linha["Fonte do chamado"]
        )

    st.markdown("---")

    # --------------------------------------------------------
    # ASSUNTO
    # --------------------------------------------------------

    st.markdown(
        "### 📝 Assunto"
    )

    st.write(
        linha["Assunto do chamado"]
    )

    # --------------------------------------------------------
    # DATAS
    # --------------------------------------------------------

    st.markdown(
        "### 🕐 Datas"
    )

    col1, col2, col3 = st.columns(3)

    with col1:

        st.write(
            "**Criação:**"
        )

        st.write(
            linha["Data de criação"]
            .strftime("%d/%m/%Y %H:%M")
            if pd.notna(
                linha["Data de criação"]
            )
            else "Não informado"
        )

    with col2:

        st.write(
            "**Última atualização:**"
        )

        st.write(
            linha["Última atualização"]
            .strftime("%d/%m/%Y %H:%M")
            if pd.notna(
                linha["Última atualização"]
            )
            else "Não informado"
        )

    with col3:

        st.write(
            "**Tempo de atendimento:**"
        )

        st.write(
            formatar_tempo(
                linha[
                    "Tempo de atendimento_min"
                ]
            )
        )

    # --------------------------------------------------------
    # HISTÓRICO
    # --------------------------------------------------------

    st.markdown(
        "### 📚 Histórico"
    )

    st.write(
        linha[
            "Histórico de Assuntos do chamado"
        ]
    )

    # --------------------------------------------------------
    # TAGS
    # --------------------------------------------------------

    st.markdown(
        "### 🏷️ Tags"
    )

    st.write(
        linha["Tags"]
    )

    # --------------------------------------------------------
    # AVALIAÇÃO
    # --------------------------------------------------------

    st.markdown(
        "### ⭐ Avaliação"
    )

    st.write(
        linha["Avaliação"]
    )

    # --------------------------------------------------------
    # COMENTÁRIO
    # --------------------------------------------------------

    st.markdown(
        "### 💬 Comentário"
    )

    st.write(
        linha["Comentário"]
    )

    st.markdown("---")

    # --------------------------------------------------------
    # EXPORTAÇÃO
    # --------------------------------------------------------

    botoes_exportacao(
        chamado,
        f"chamado_{numero}"
    )
# ============================================================
# PENDÊNCIAS
# ============================================================

def tela_pendencias(df):

    st.markdown(
        '<div class="titulo-secao">⏳ PENDÊNCIAS</div>',
        unsafe_allow_html=True
    )

    pendencias = df[
        df["Status do chamado"]
        .isin(STATUS_PENDENTES)
    ].copy()

    if pendencias.empty:

        st.success(
            "Não existem pendências com os filtros selecionados."
        )

        return

    total = len(pendencias)

    sem_atendente = len(
        pendencias[
            pendencias["Atendente"]
            == "Não atribuído"
        ]
    )

    tempo_medio = (
        pendencias[
            "Tempo de atendimento_min"
        ].mean()
    )

    col1, col2, col3 = st.columns(3)

    with col1:
        card(
            "Total pendente",
            total
        )

    with col2:
        card(
            "Sem atendente",
            sem_atendente
        )

    with col3:
        card(
            "Tempo médio",
            formatar_tempo(
                tempo_medio
            )
        )

    st.markdown("---")

    tabela = preparar_exportacao(
        pendencias
    )

    st.dataframe(
        tabela,
        use_container_width=True,
        hide_index=True
    )

    st.markdown(
        "### 📥 Exportação"
    )

    botoes_exportacao(
        pendencias,
        "pendencias_layers"
    )


# ============================================================
# ANÁLISES
# ============================================================

def tela_analises(df):

    st.markdown(
        '<div class="titulo-secao">📈 ANÁLISES</div>',
        unsafe_allow_html=True
    )

    if df.empty:

        st.warning(
            "Nenhum dado disponível para análise."
        )

        return

    # --------------------------------------------------------
    # 1. Atendimentos por atendente
    # --------------------------------------------------------

    st.markdown(
        "### 1. Atendimentos por atendente"
    )

    tabela = tabela_quantitativa(
        df,
        "Atendente"
    )

    st.dataframe(
        tabela,
        use_container_width=True,
        hide_index=True
    )

    fig = grafico_atendentes(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 2. Tempo médio por atendente
    # --------------------------------------------------------

    st.markdown(
        "### 2. Tempo médio por atendente"
    )

    if not df.empty:

        tabela_tempo = (
            df.groupby(
                "Atendente"
            )["Tempo de atendimento_min"]
            .agg(
                [
                    "count",
                    "mean",
                    "median"
                ]
            )
            .reset_index()
        )

        tabela_tempo.columns = [
            "Atendente",
            "Quantidade",
            "Tempo médio (min)",
            "Mediana (min)"
        ]

        tabela_tempo[
            "Tempo médio"
        ] = tabela_tempo[
            "Tempo médio (min)"
        ].apply(
            formatar_tempo
        )

        tabela_tempo[
            "Mediana"
        ] = tabela_tempo[
            "Mediana (min)"
        ].apply(
            formatar_tempo
        )

        st.dataframe(
            tabela_tempo[
                [
                    "Atendente",
                    "Quantidade",
                    "Tempo médio",
                    "Mediana"
                ]
            ],
            use_container_width=True,
            hide_index=True
        )

    fig = grafico_tempo_atendente(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 3. Status por atendente
    # --------------------------------------------------------

    st.markdown(
        "### 3. Status por atendente"
    )

    fig = grafico_status_atendente(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 4. Chamados por fonte
    # --------------------------------------------------------

    st.markdown(
        "### 4. Chamados por fonte"
    )

    tabela = tabela_quantitativa(
        df,
        "Fonte do chamado",
        "Fonte"
    )

    st.dataframe(
        tabela,
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # 5. Tempo médio por fonte
    # --------------------------------------------------------

    st.markdown(
        "### 5. Tempo médio por fonte"
    )

    fig = grafico_tempo_fonte(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 6. Chamados por canal
    # --------------------------------------------------------

    st.markdown(
        "### 6. Chamados por canal"
    )

    tabela = tabela_quantitativa(
        df,
        "Nome do canal",
        "Canal"
    )

    st.dataframe(
        tabela,
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # 7. Status por canal
    # --------------------------------------------------------

    st.markdown(
        "### 7. Status por canal"
    )

    fig = grafico_status_canal(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 8. Evolução mensal
    # --------------------------------------------------------

    st.markdown(
        "### 8. Evolução mensal"
    )

    fig = grafico_mensal(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 9. Dia da semana
    # --------------------------------------------------------

    st.markdown(
        "### 9. Chamados por dia da semana"
    )

    fig = grafico_dia_semana(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 10. Hora do dia
    # --------------------------------------------------------

    st.markdown(
        "### 10. Chamados por horário"
    )

    fig = grafico_hora(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 11. Principais assuntos
    # --------------------------------------------------------

    st.markdown(
        "### 11. Principais assuntos"
    )

    tabela = tabela_quantitativa(
        df,
        "Assunto do chamado",
        "Assunto"
    )

    st.dataframe(
        tabela.head(20),
        use_container_width=True,
        hide_index=True
    )

    fig = grafico_assuntos(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 12. Taxa de resolução
    # --------------------------------------------------------

    st.markdown(
        "### 12. Taxa de resolução"
    )

    total = len(df)

    resolvidos = len(
        df[
            df["Status do chamado"]
            == "Resolvido"
        ]
    )

    taxa = (
        resolvidos / total * 100
        if total > 0
        else 0
    )

    col1, col2 = st.columns(2)

    with col1:

        card(
            "Chamados resolvidos",
            resolvidos
        )

    with col2:

        card(
            "Taxa de resolução",
            f"{taxa:.1f}%"
        )

    # --------------------------------------------------------
    # 13. Envelhecimento das pendências
    # --------------------------------------------------------

    st.markdown(
        "### 13. Envelhecimento das pendências"
    )

    pendencias = df[
        df["Status do chamado"]
        .isin(STATUS_PENDENTES)
    ].copy()

    if not pendencias.empty:

        agora = pd.Timestamp.now()

        pendencias[
            "Dias em aberto"
        ] = (
            agora
            - pendencias[
                "Data de criação"
            ]
        ).dt.total_seconds() / 86400

        def classificar_idade(dias):

            if pd.isna(dias):
                return "Não informado"

            if dias <= 1:
                return "Até 1 dia"

            if dias <= 3:
                return "2 a 3 dias"

            if dias <= 7:
                return "4 a 7 dias"

            if dias <= 15:
                return "8 a 15 dias"

            return "Mais de 15 dias"

        pendencias[
            "Faixa"
        ] = pendencias[
            "Dias em aberto"
        ].apply(
            classificar_idade
        )

        ordem = [
            "Até 1 dia",
            "2 a 3 dias",
            "4 a 7 dias",
            "8 a 15 dias",
            "Mais de 15 dias"
        ]

        envelhecimento = (
            pendencias["Faixa"]
            .value_counts()
            .reindex(
                ordem,
                fill_value=0
            )
            .reset_index()
        )

        envelhecimento.columns = [
            "Faixa",
            "Quantidade"
        ]

        st.dataframe(
            envelhecimento,
            use_container_width=True,
            hide_index=True
        )

        fig = px.bar(
            envelhecimento,
            x="Faixa",
            y="Quantidade",
            text="Quantidade",
            title="Envelhecimento das pendências"
        )

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # 14. Avaliações
    # --------------------------------------------------------

    st.markdown(
        "### 14. Avaliações dos atendimentos"
    )

    avaliacao = df[
        df["Avaliação"]
        != "Não informado"
    ].copy()

    if avaliacao.empty:

        st.info(
            "Não existem avaliações informadas."
        )

    else:

        tabela = tabela_quantitativa(
            avaliacao,
            "Avaliação"
        )

        st.dataframe(
            tabela,
            use_container_width=True,
            hide_index=True
        )

    # --------------------------------------------------------
    # 15. Principais solicitantes
    # --------------------------------------------------------

    st.markdown(
        "### 15. Principais solicitantes"
    )

    tabela = tabela_quantitativa(
        df,
        "Nome do solicitante",
        "Solicitante"
    )

    st.dataframe(
        tabela.head(20),
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # Exportação
    # --------------------------------------------------------

    st.markdown("---")

    st.markdown(
        "### 📥 Exportação das análises"
    )

    botoes_exportacao(
        df,
        "analises_layers"
    )


# ============================================================
# RELATÓRIO GERAL
# ============================================================

def tela_relatorio_geral(df):

    st.markdown(
        '<div class="titulo-secao">📑 RELATÓRIO GERAL</div>',
        unsafe_allow_html=True
    )

    if df.empty:

        st.warning(
            "Nenhum dado disponível para o relatório."
        )

        return

    # --------------------------------------------------------
    # Resumo
    # --------------------------------------------------------

    total = len(df)

    novos = len(
        df[
            df["Status do chamado"]
            == "Novo"
        ]
    )

    abertos = len(
        df[
            df["Status do chamado"]
            == "Aberto"
        ]
    )

    aguardando = len(
        df[
            df["Status do chamado"]
            == "Aguardando"
        ]
    )

    resolvidos = len(
        df[
            df["Status do chamado"]
            == "Resolvido"
        ]
    )

    pendentes = len(
        df[
            df["Status do chamado"]
            .isin(STATUS_PENDENTES)
        ]
    )

    col1, col2, col3, col4, col5, col6 = st.columns(6)

    with col1:
        card("Total", total)

    with col2:
        card("Novos", novos)

    with col3:
        card("Abertos", abertos)

    with col4:
        card("Aguardando", aguardando)

    with col5:
        card("Resolvidos", resolvidos)

    with col6:
        card("Pendentes", pendentes)

    st.markdown("---")

    # --------------------------------------------------------
    # Status
    # --------------------------------------------------------

    st.markdown(
        "### Chamados por status"
    )

    fig = grafico_status(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    st.dataframe(
        tabela_quantitativa(
            df,
            "Status do chamado",
            "Status"
        ),
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # Canais
    # --------------------------------------------------------

    st.markdown(
        "### Chamados por canal"
    )

    fig = grafico_canais(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    st.dataframe(
        tabela_quantitativa(
            df,
            "Nome do canal",
            "Canal"
        ),
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # Atendentes
    # --------------------------------------------------------

    st.markdown(
        "### Chamados por atendente"
    )

    fig = grafico_atendentes(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    st.dataframe(
        tabela_quantitativa(
            df,
            "Atendente"
        ),
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # Assuntos
    # --------------------------------------------------------

    st.markdown(
        "### Principais assuntos"
    )

    fig = grafico_assuntos(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    st.dataframe(
        tabela_quantitativa(
            df,
            "Assunto do chamado",
            "Assunto"
        ).head(20),
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # Evolução
    # --------------------------------------------------------

    st.markdown(
        "### Evolução mensal"
    )

    fig = grafico_mensal(df)

    if fig:

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    # --------------------------------------------------------
    # Pendências
    # --------------------------------------------------------

    st.markdown(
        "### Pendências"
    )

    pendencias = df[
        df["Status do chamado"]
        .isin(STATUS_PENDENTES)
    ]

    st.dataframe(
        preparar_exportacao(
            pendencias
        ),
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # Exportação geral
    # --------------------------------------------------------

    st.markdown("---")

    botoes_relatorio_geral(
        df
    )


# ============================================================
# MAIN
# ============================================================

def main():

    # --------------------------------------------------------
    # LOGIN
    # --------------------------------------------------------

    if not st.session_state.get(
        "autenticado",
        False
    ):

        tela_login()

        return

    # --------------------------------------------------------
    # CABEÇALHO
    # --------------------------------------------------------

    st.markdown(
        '<div class="titulo-principal">'
        'SGA Layers'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="subtitulo">'
        'Sistema de Gestão de Atendimentos Layers'
        '</div>',
        unsafe_allow_html=True
    )

    # --------------------------------------------------------
    # SIDEBAR
    # --------------------------------------------------------

    st.sidebar.markdown(
        "## 📊 SGA Layers"
    )

    st.sidebar.write(
        f"Usuário: "
        f"**{st.session_state.get('usuario', '')}**"
    )

    if st.sidebar.button(
        "🚪 Sair",
        use_container_width=True
    ):

        st.session_state[
            "autenticado"
        ] = False

        st.session_state.pop(
            "usuario",
            None
        )

        st.rerun()

    st.sidebar.markdown("---")

    pagina = st.sidebar.radio(
        "Navegação",
        [
            "📊 Dashboard",
            "📋 Chamados",
            "🔍 Detalhamento",
            "📈 Análises",
            "⏳ Pendências",
            "📑 Relatório Geral"
        ]
    )

    st.sidebar.markdown("---")

    # --------------------------------------------------------
    # UPLOAD
    # --------------------------------------------------------

    arquivo = st.sidebar.file_uploader(
        "📁 Carregar relatório CSV",
        type=["csv"]
    )

    if arquivo is None:

        st.info(
            "👈 Carregue o relatório CSV do Layers "
            "para iniciar a análise."
        )

        return

    # --------------------------------------------------------
    # LEITURA
    # --------------------------------------------------------

    try:

        df = ler_csv(
            arquivo
        )

    except Exception as erro:

        st.error(
            f"Erro ao ler o CSV: {erro}"
        )

        return

    # --------------------------------------------------------
    # VALIDAÇÃO
    # --------------------------------------------------------

    colunas_faltantes = [
        coluna
        for coluna in COLUNAS_ESPERADAS
        if coluna not in df.columns
    ]

    if colunas_faltantes:

        st.error(
            "O arquivo não possui todas as colunas esperadas."
        )

        st.write(
            "Colunas faltantes:"
        )

        st.write(
            colunas_faltantes
        )

        st.write(
            "Colunas encontradas:"
        )

        st.write(
            df.columns.tolist()
        )

        return

    # --------------------------------------------------------
    # TRATAMENTO
    # --------------------------------------------------------

    df = tratar_dados(
        df
    )

    # --------------------------------------------------------
    # FILTROS GLOBAIS
    # --------------------------------------------------------

    df_filtrado = aplicar_filtros_globais(
        df
    )

    # --------------------------------------------------------
    # INFORMAÇÃO
    # --------------------------------------------------------

    st.sidebar.markdown("---")

    st.sidebar.write(
        f"📁 Total original: **{len(df)}**"
    )

    st.sidebar.write(
        f"🔎 Após filtros: **{len(df_filtrado)}**"
    )

    # --------------------------------------------------------
    # PÁGINAS
    # --------------------------------------------------------

    if pagina == "📊 Dashboard":

        tela_dashboard(
            df_filtrado
        )

    elif pagina == "📋 Chamados":

        tela_chamados(
            df_filtrado
        )

    elif pagina == "🔍 Detalhamento":

        tela_detalhamento(
            df_filtrado
        )

    elif pagina == "📈 Análises":

        tela_analises(
            df_filtrado
        )

    elif pagina == "⏳ Pendências":

        tela_pendencias(
            df_filtrado
        )

    elif pagina == "📑 Relatório Geral":

        tela_relatorio_geral(
            df_filtrado
        )


# ============================================================
# EXECUÇÃO
# ============================================================

if __name__ == "__main__":
    main()
